#!/usr/bin/env python3
"""
Full extraction of every file in data/New-OA-Data/.
Outputs: extracted/ folder with per-file artifacts + FULL_EXTRACTION_REPORT.md
Requires: pyreadstat, pandas, openpyxl, python-docx, pypdf (pip install pyreadstat pandas openpyxl python-docx pypdf)
"""

import sys
from pathlib import Path
import json

BASE = Path(__file__).resolve().parents[1]
DATA_ROOT = BASE / "data" / "New-OA-Data"
OUT = DATA_ROOT / "extracted"
OUT.mkdir(parents=True, exist_ok=True)

def log(msg):
    print(msg)
    sys.stdout.flush()

# ---------------------------------------------------------------------------
# 1. SPSS .sav files – full column list + labels + row count
# ---------------------------------------------------------------------------
def extract_sav(path: Path) -> dict:
    try:
        import pyreadstat
        df, meta = pyreadstat.read_sav(str(path))
    except Exception as e:
        return {"ok": False, "error": str(e), "path": str(path)}
    cols = list(df.columns)
    labels = {}
    try:
        if meta is not None and getattr(meta, "column_names_to_labels", None):
            # Copy to plain dict to avoid C-extension quirks; limit size
            raw = meta.column_names_to_labels
            keys = list(raw.keys())[:5000]
            for k in keys:
                v = raw.get(k)
                labels[str(k)] = str(v) if v is not None else ""
    except Exception:
        pass
    dtypes = {}
    for c in cols[:100]:
        try:
            dtypes[c] = str(df[c].dtype)
        except Exception:
            dtypes[c] = "unknown"
    return {
        "ok": True,
        "path": str(path),
        "rows": len(df),
        "columns": cols,
        "column_count": len(cols),
        "labels": labels,
        "dtypes": dtypes,
    }

# ---------------------------------------------------------------------------
# 2. SPSS .por files
# ---------------------------------------------------------------------------
def extract_por(path: Path) -> dict:
    try:
        import pyreadstat
        df, meta = pyreadstat.read_por(str(path))
        cols = list(df.columns)
        labels = {}
        if meta.column_names_to_labels:
            labels = {k: str(v) for k, v in meta.column_names_to_labels.items()}
        return {
            "ok": True,
            "path": str(path),
            "rows": len(df),
            "columns": cols,
            "column_count": len(cols),
            "labels": labels,
        }
    except Exception as e:
        return {"ok": False, "error": str(e), "path": str(path)}

# ---------------------------------------------------------------------------
# 3. Excel .xlsx – every sheet, full content (or row count + column names)
# ---------------------------------------------------------------------------
def extract_xlsx(path: Path) -> dict:
    try:
        import pandas as pd
        xl = pd.ExcelFile(path)
        sheets = {}
        for name in xl.sheet_names:
            df = pd.read_excel(xl, sheet_name=name, header=None)
            nrows, ncols = len(df), len(df.columns)
            # Store full content for writing CSV; avoid keeping huge lists in JSON
            full_content = df.fillna("").astype(str).values.tolist()
            sheets[name] = {
                "rows": nrows,
                "cols": ncols,
                "columns_first_row": df.iloc[0].astype(str).tolist() if nrows > 0 else [],
                "full_content": full_content,
            }
        return {"ok": True, "path": str(path), "sheet_names": xl.sheet_names, "sheets": sheets}
    except Exception as e:
        return {"ok": False, "error": str(e), "path": str(path)}

# ---------------------------------------------------------------------------
# 4. PDF – full text
# ---------------------------------------------------------------------------
def extract_pdf(path: Path) -> dict:
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        text_by_page = []
        for i, page in enumerate(reader.pages):
            t = page.extract_text()
            text_by_page.append({"page": i + 1, "text": t or ""})
        full_text = "\n\n".join(p["text"] for p in text_by_page)
        return {
            "ok": True,
            "path": str(path),
            "num_pages": len(reader.pages),
            "text_by_page": text_by_page,
            "full_text": full_text,
        }
    except Exception as e:
        return {"ok": False, "error": str(e), "path": str(path)}

# ---------------------------------------------------------------------------
# 5. Word .docx – full text
# ---------------------------------------------------------------------------
def extract_docx(path: Path) -> dict:
    try:
        from docx import Document
        doc = Document(path)
        paras = [p.text for p in doc.paragraphs]
        full_text = "\n".join(paras)
        tables_text = []
        for t in doc.tables:
            for row in t.rows:
                cells = [c.text for c in row.cells]
                tables_text.append(" | ".join(cells))
        if tables_text:
            full_text += "\n\n[TABLES]\n" + "\n".join(tables_text)
        return {
            "ok": True,
            "path": str(path),
            "num_paragraphs": len(doc.paragraphs),
            "num_tables": len(doc.tables),
            "full_text": full_text,
        }
    except Exception as e:
        return {"ok": False, "error": str(e), "path": str(path)}

# ---------------------------------------------------------------------------
# Run extraction on all files
# ---------------------------------------------------------------------------
def main():
    log("=== FULL EXTRACTION: data/New-OA-Data ===")

    # Gather all files
    check_dir = DATA_ROOT / "CHECK"
    klsa_dir = DATA_ROOT / "KLoSa"
    klsa_spss = klsa_dir / "KLoSA 1-9th wave (SPSS)"

    files_sav = []
    if (check_dir / "CHECK_T0_DANS_nsin_ENG_20161128.sav").exists():
        files_sav.append(check_dir / "CHECK_T0_DANS_nsin_ENG_20161128.sav")
    if klsa_spss.exists():
        files_sav.extend(sorted(klsa_spss.glob("*.sav")))

    files_por = []
    if (check_dir / "CHECK_Radiographic_Scoring_OA_T0T2T5T8T10_DANS_nsinENG_20170726.por").exists():
        files_por.append(check_dir / "CHECK_Radiographic_Scoring_OA_T0T2T5T8T10_DANS_nsinENG_20170726.por")

    files_xlsx = list(klsa_dir.glob("*.xlsx")) if klsa_dir.exists() else []

    files_pdf = list(check_dir.glob("*.pdf")) if check_dir.exists() else []
    files_docx = list(klsa_dir.glob("*.docx")) if klsa_dir.exists() else []

    results = {"sav": {}, "por": {}, "xlsx": {}, "pdf": {}, "docx": {}}

    # SAV
    log(f"\n--- .sav ({len(files_sav)} files) ---")
    for p in files_sav:
        rel = p.relative_to(DATA_ROOT)
        log(f"  {rel}")
        r = extract_sav(p)
        results["sav"][str(rel)] = r
        if r.get("ok"):
            out_name = p.stem + "_sav_columns.json"
            out_path = OUT / "sav"
            out_path.mkdir(exist_ok=True)
            with open(out_path / out_name, "w", encoding="utf-8") as f:
                json.dump({"rows": r["rows"], "columns": r["columns"], "labels": r["labels"]}, f, indent=2)
        else:
            log(f"    ERROR: {r.get('error')}")

    # POR
    log(f"\n--- .por ({len(files_por)} files) ---")
    for p in files_por:
        rel = p.relative_to(DATA_ROOT)
        log(f"  {rel}")
        r = extract_por(p)
        results["por"][str(rel)] = r
        if not r.get("ok"):
            log(f"    ERROR: {r.get('error')}")

    # XLSX
    log(f"\n--- .xlsx ({len(files_xlsx)} files) ---")
    for p in files_xlsx:
        rel = p.relative_to(DATA_ROOT)
        log(f"  {rel}")
        r = extract_xlsx(p)
        results["xlsx"][str(rel)] = {"ok": r.get("ok"), "path": r.get("path"), "error": r.get("error"), "sheet_names": r.get("sheet_names"), "sheet_row_cols": {}}
        if r.get("ok"):
            for sh, data in r.get("sheets", {}).items():
                results["xlsx"][str(rel)]["sheet_row_cols"][sh] = {"rows": data["rows"], "cols": data["cols"]}
                out_dir = OUT / "xlsx"
                out_dir.mkdir(exist_ok=True)
                safe_name = p.stem.replace(" ", "_").replace("~", "_")
                sheet_file = out_dir / f"{safe_name}_{sh.replace(' ', '_')}.csv"
                try:
                    import pandas as pd
                    df = pd.DataFrame(data["full_content"])
                    df.to_csv(sheet_file, index=False, header=False)
                except Exception:
                    pass
        else:
            log(f"    ERROR: {r.get('error')}")

    # PDF
    log(f"\n--- .pdf ({len(files_pdf)} files) ---")
    for p in files_pdf:
        rel = p.relative_to(DATA_ROOT)
        log(f"  {rel}")
        r = extract_pdf(p)
        results["pdf"][str(rel)] = {"ok": r.get("ok"), "path": r.get("path"), "error": r.get("error"), "num_pages": r.get("num_pages")}
        if r.get("ok"):
            out_dir = OUT / "pdf"
            out_dir.mkdir(exist_ok=True)
            txt_path = out_dir / (p.stem + "_full_text.txt")
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(r.get("full_text", ""))
            results["pdf"][str(rel)]["full_text_preview"] = (r.get("full_text") or "")[:5000]
        else:
            log(f"    ERROR: {r.get('error')}")

    # DOCX
    log(f"\n--- .docx ({len(files_docx)} files) ---")
    for p in files_docx:
        rel = p.relative_to(DATA_ROOT)
        log(f"  {rel}")
        r = extract_docx(p)
        results["docx"][str(rel)] = {"ok": r.get("ok"), "path": r.get("path"), "error": r.get("error"), "num_paragraphs": r.get("num_paragraphs"), "num_tables": r.get("num_tables")}
        if r.get("ok"):
            out_dir = OUT / "docx"
            out_dir.mkdir(exist_ok=True)
            txt_path = out_dir / (p.stem + "_full_text.txt")
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(r.get("full_text", ""))
            results["docx"][str(rel)]["full_text_preview"] = (r.get("full_text") or "")[:5000]
        else:
            log(f"    ERROR: {r.get('error')}")

    # Save summary JSON
    summary = {
        "sav": {k: {"ok": v.get("ok"), "rows": v.get("rows"), "columns": v.get("column_count")} for k, v in results["sav"].items()},
        "por": {k: {"ok": v.get("ok"), "error": v.get("error")} for k, v in results["por"].items()},
        "xlsx": {k: {"ok": v.get("ok"), "sheets": v.get("sheet_row_cols")} for k, v in results["xlsx"].items()},
        "pdf": {k: {"ok": v.get("ok"), "num_pages": v.get("num_pages")} for k, v in results["pdf"].items()},
        "docx": {k: {"ok": v.get("ok"), "num_paragraphs": v.get("num_paragraphs")} for k, v in results["docx"].items()},
    }
    with open(OUT / "extraction_summary.json", "w", encoding="utf-8") as f:
        json.dump(summary, f, indent=2)

    # Write FULL_EXTRACTION_REPORT.md
    report_lines = [
        "# Full Extraction Report: data/New-OA-Data",
        "",
        "Every file in New-OA-Data was processed. Artifacts are in `extracted/`.",
        "",
        "## 1. SPSS .sav files",
        ""
    ]
    for k, v in results["sav"].items():
        if v.get("ok"):
            report_lines.append(f"- **{k}**: {v['rows']} rows × {v['column_count']} columns. Columns/labels → `extracted/sav/{Path(k).stem}_sav_columns.json`")
        else:
            report_lines.append(f"- **{k}**: ERROR – {v.get('error', '')}")
        report_lines.append("")
    report_lines.extend(["## 2. SPSS .por files", ""])
    for k, v in results["por"].items():
        if v.get("ok"):
            report_lines.append(f"- **{k}**: {v['rows']} rows × {v['column_count']} columns")
        else:
            report_lines.append(f"- **{k}**: ERROR – {v.get('error', '')}")
        report_lines.append("")
    report_lines.extend(["## 3. Excel .xlsx files", ""])
    for k, v in results["xlsx"].items():
        if v.get("ok"):
            report_lines.append(f"- **{k}**: sheets {v.get('sheet_names', [])}. Per-sheet CSV → `extracted/xlsx/`")
        else:
            report_lines.append(f"- **{k}**: ERROR – {v.get('error', '')}")
        report_lines.append("")
    report_lines.extend(["## 4. PDF files", ""])
    for k, v in results["pdf"].items():
        if v.get("ok"):
            report_lines.append(f"- **{k}**: {v.get('num_pages')} pages. Full text → `extracted/pdf/{Path(k).stem}_full_text.txt`")
        else:
            report_lines.append(f"- **{k}**: ERROR – {v.get('error', '')}")
        report_lines.append("")
    report_lines.extend(["## 5. Word .docx files", ""])
    for k, v in results["docx"].items():
        if v.get("ok"):
            report_lines.append(f"- **{k}**: {v.get('num_paragraphs')} paragraphs, {v.get('num_tables')} tables. Full text → `extracted/docx/{Path(k).stem}_full_text.txt`")
        else:
            report_lines.append(f"- **{k}**: ERROR – {v.get('error', '')}")
        report_lines.append("")
    with open(OUT / "FULL_EXTRACTION_REPORT.md", "w", encoding="utf-8") as f:
        f.write("\n".join(report_lines))

    log(f"\nDone. Outputs in {OUT}")
    return results

if __name__ == "__main__":
    main()
